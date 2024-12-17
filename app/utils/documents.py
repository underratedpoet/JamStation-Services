from docx import Document

def save_report_to_docx(equipment_data, checks_data, stats_data, comments, file_path):
    doc = Document()
    doc.add_heading('Отчет о состоянии оборудования', 0)

    doc.add_heading('Оборудование:', level=1)
    doc.add_paragraph(equipment_data)

    doc.add_heading('Проверки:', level=1)
    doc.add_paragraph(checks_data)

    doc.add_heading('Сводная статистика:', level=1)
    doc.add_paragraph(stats_data)

    doc.add_heading('Комментарии:', level=1)
    doc.add_paragraph(comments)

    doc.save(file_path)

def save_standart_report_to_docx(comments, file_path, emp_id):
    doc = Document()
    doc.add_heading('Стандвртный отчет', 0)

    doc.add_heading(f'Сотрудник ID: {emp_id}', level=1)
    doc.add_paragraph(comments)

    doc.save(file_path)    